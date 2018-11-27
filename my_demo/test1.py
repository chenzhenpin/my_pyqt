# -*- coding: utf-8 -*-

"""
Module implementing MainWindow.
"""

from PyQt5.QtCore import pyqtSlot
from PyQt5.QtWidgets import QMainWindow
from PyQt5.QtCore import *
from Ui_test1 import Ui_MainWindow
from PyQt5 import QtWidgets,QtGui
from PyQt5 import *

from PyQt5.QtWidgets import *
import webbrowser,time


class MainWindow(QMainWindow, Ui_MainWindow):
    """
    Class documentation goes here.
    """

    def __init__(self, parent=None):
        """
        Constructor

        @param parent reference to the parent widget
        @type QWidget
        """
        super(MainWindow, self).__init__(parent)
        self.setupUi(self)
        self.graphicsView.mousePressEvent = self.img_clicked

    def img_clicked(self,e):
        webbrowser.open('http://www.baidu.com')
        print('ok')

    @pyqtSlot()
    def on_pushButton_4_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        # raise NotImplementedError
        self.label_13.setText('我去，你居然敢按我')

    @pyqtSlot()
    def on_pushButton_5_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        self.lineEdit.setText('')

    @pyqtSlot()
    def on_pushButton_6_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        str = self.lineEdit.text()
        self.textBrowser.append(str)
        print(self.textBrowser.toPlainText())

    @pyqtSlot()
    def on_radioButton_6_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        self.label_5.setText('弃权')

    @pyqtSlot()
    def on_radioButton_7_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        self.radioButton_5.setChecked(True)

    @pyqtSlot()
    def on_radioButton_8_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet

        self.radioButton_4.setChecked(True)

    @pyqtSlot()
    def on_radioButton_4_clicked(self):
        """
        Slot documentation goes here.
        """
        if self.radioButton_7.isChecked():
            self.label_5.setText('你不同意')
            print('你不同意')
        elif self.radioButton_8.isChecked():
            self.label_5.setText('你同意')
            print('你同意')
        else:
            self.label_5.setText('你选择弃权')
            print('你选择弃权')

    @pyqtSlot()
    def on_radioButton_5_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        print('你取消')

    @pyqtSlot(int)
    def on_horizontalSlider_valueChanged(self, value):
        """
        Slot documentation goes here.

        @param value DESCRIPTION
        @type int
        """
        # TODO: not implemented yet
        self.progressBar.setProperty("value", value)

    @pyqtSlot(int)
    def on_dial_valueChanged(self, value):
        """
        Slot documentation goes here.

        @param value DESCRIPTION
        @type int
        """
        # TODO: not implemented yet
        self.lcdNumber.display(value)

    @pyqtSlot()
    def on_pushButton_7_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        my_btn = QMessageBox.information(self, '信息框', '信息')
        print(my_btn)

    @pyqtSlot()
    def on_pushButton_8_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        my_btn = QMessageBox.question(self, u'对话框', u'是否保存', QMessageBox.Yes, QMessageBox.No)
        print(my_btn)

    @pyqtSlot()
    def on_pushButton_9_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        my_btn = QMessageBox.warning(self, '警告框', '你的信息错误！', QMessageBox.Ok, QMessageBox.Cancel)
        print(my_btn)

    @pyqtSlot()
    def on_pushButton_10_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        my_btn = QMessageBox.critical(self, '错误框', '你的信息错误！', QMessageBox.Ok, QMessageBox.Cancel)
        print(my_btn)

    @pyqtSlot()
    def on_pushButton_11_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        my_btn = QMessageBox.about(self, '关于框', '你的信息错误！')
        my_btn = QMessageBox.aboutQt(self, 'aboutQt')
        print(my_btn)

    @pyqtSlot()
    def on_pushButton_12_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        my_str, ok = QInputDialog.getText(self, '请输入字符', '在此输入', QLineEdit.Normal, '请在此输入')
        print(my_str)

    @pyqtSlot()
    def on_pushButton_13_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        my_str, ok = QInputDialog.getInt(self, '请输入整数', '在此输入', 30, 0, 100)
        print(my_str)

    @pyqtSlot()
    def on_pushButton_14_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented
        my_list = []
        my_list.append('很不错')
        my_list.append('很好')
        my_str, ok = QInputDialog.getItem(self, '请选择下拉框', '选项', my_list, QMessageBox.Ok, QMessageBox.Cancel)
        print(my_str)

    @pyqtSlot()
    def on_pushButton_15_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        from info import Dialog
        d = Dialog()
        d.exec_()
        
        
    @pyqtSlot()
    def on_action_triggered(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        print('创建')

    @pyqtSlot()
    def on_action_5_triggered(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        print('打开')
        

        
        
    @pyqtSlot()
    def on_actionDakai_triggered(self):
        """
        Slot documentation goes here.
        """
        # TO

        my_file_path=QFileDialog.getOpenFileName(self,'打开文件','/')[0]
        print(my_file_path)
        if my_file_path[-4:] =='.doc' or my_file_path[-5:] =='.docx':
            # from win32com import client as wc
            # word=wc.Dispatch('Word.Application')
            # word.Visible=1
            # my_word=word.Documents.Open(my_file_path.replace('/','\\'))
            # my_count=my_word.Paragraphs.Count
            # for i in range(my_count):
            #     my_pr=my_word.Paragraphs[i].Range
            #     print(my_pr)
            #     time.sleep(3)
            # my_word.Close()
            import docx
            doc=docx.Document(my_file_path.replace('/','\\'))
            # data=doc.paragraphs[0].text
            # print(data)
            for my_paragraphs in doc.paragraphs:
                self.textBrowser.append(my_paragraphs.text)
        elif my_file_path[-4:] =='.txt':
            f=open(my_file_path)
            my_data=f.read()
            print(my_data)
            f.close()

if __name__ == "__main__":
    import sys

    app = QtWidgets.QApplication(sys.argv)
    ui = MainWindow()
    splash=QSplashScreen(QtGui.QPixmap(":/frx/0.jpg"))
    splash.show()
    splash.showMessage('正在启动',Qt.AlignVCenter|Qt.AlignHCenter,Qt.black)
    time.sleep(1)
    ui.show()
    splash.close()
    sys.exit(app.exec_())
    
